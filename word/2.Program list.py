from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn

doc2 = Document()

# Insert a Unordered List
doc2.add_paragraph('下列哪个是动物？')

doc2.add_paragraph('Apple', style='List Bullet')
doc2.add_paragraph('Happiness', style='List Bullet')
doc2.add_paragraph('Laziness', style='List Bullet')
doc2.add_paragraph('Anger', style='List Bullet')
doc2.add_paragraph('Wolf', style='List Bullet')

# Insert a Ordered List
doc2.add_paragraph('2020 annual procedure')

doc2.add_paragraph('Be a expert', style='List Number')
doc2.add_paragraph('Be happy', style='List Number')
doc2.add_paragraph('Enjoy myself', style='List Number')

# Insert Pictures
doc2.add_picture('./images/pic1.jfif', width=Inches(5.5))

# Insert Tables
doc2.add_heading('Table', 2)
# This is Header of the Table
table = doc2.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Numbering'
hdr_cells[1].text = 'Name'
hdr_cells[2].text = 'Gender'
hdr_cells[3].text = 'Height'
# This is Content of the Table
records = (
    (1, 'Lisa', 'female', 1.6),
    (2, 'Sun', 'male', 1.8),
    (3, 'Ming', 'male', 1.7),
    (4, 'Yang', 'female', 1.7)
)
# In the cell of the table, only the str is allowed
for (id, name, gender, height) in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(id)
    row_cells[1].text = name
    row_cells[2].text = gender
    row_cells[3].text = str(height)

# Manually add a Page
doc2.add_page_break()


doc2.save('word2.docx')
