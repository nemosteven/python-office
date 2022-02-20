from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import RGBColor
from docx.oxml.ns import qn


# Similar to the Excel, we need to establish a Document() first
doc1 = Document()

# Add a Title
doc1.add_heading('Title of the Document', 0)

# Paragraphs
doc1.add_paragraph('Paragraph 1')

# First-level Title
doc1.add_heading('First-level Title', 1)

# Second-level Title
doc1.add_heading('Second-level Title', 2)

# Third-level Title
doc1.add_heading('Third-level Title', 3)

# Font-Style and Reference

# This is Font
# Here add_run is specificly used to add content to the paragraph
run = doc1.add_paragraph('Here is English Font:').add_run('This Font is Times New Roman')
run.font.name = 'Times New Roman'
run = doc1.add_paragraph('Here is Chinese Font:').add_run('当前字体是黑体')
run.font.name = '黑体'

# This is Size
run = doc1.add_paragraph('Here is 20 font size:').add_run('Font Size 20')
run.font.size = Pt(20)

# This is italics
run = doc1.add_paragraph('Here is italics:').add_run('Italics')
run.italic = True

# This is under-lining
doc1.add_paragraph('Here is under-lining:').add_run('Under-lining').underline = True

# This is bold
doc1.add_paragraph('Here is bold:').add_run('Bold').bold = True

# This is Font-Color Red
run = doc1.add_paragraph('Here is red:').add_run('Red')
run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

# Here is the Reference
doc1.add_paragraph('Here is the Reference', style='Intense Quote')

# Save the document
doc1.save('word1.docx')