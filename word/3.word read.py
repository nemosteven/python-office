from docx import Document

# Open the document to read
doc1 = Document('word1.docx')

pl = [paragraph.text for paragraph in doc1.paragraphs]

# Output the document 1
print('####### Output of word1.docx #######')
for i in pl:
    print(i)

print()

# Output the document 2
doc2 = Document('word2.docx')
print('####### Output of word2.docx #######')
for para in doc2.paragraphs:
    print(para.text)
print('Attention: Table and Picture are unable to display')

# Print the Table
for table in doc2.tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text, end=' ')
        print()
    print()

